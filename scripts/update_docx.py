
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
MD_SOURCE = "/Users/musfiqurtuhin/Documents/WorkSpace/political_violence_tracker/docs/PROJECT_MASTER_DOCUMENTATION.md"
DOCX_TARGET = "/Users/musfiqurtuhin/Documents/WorkSpace/political_violence_tracker/https___violencetracker.org_.docx"

def read_markdown(path):
    with open(path, 'r', encoding='utf-8') as f:
        return f.readlines()

def update_docx():
    print(f"Reading Markdown from: {MD_SOURCE}")
    lines = read_markdown(MD_SOURCE)
    
    print(f"Opening Word Doc: {DOCX_TARGET}")
    try:
        doc = Document(DOCX_TARGET)
    except Exception as e:
        print(f"Could not open existing doc, creating new one. Error: {e}")
        doc = Document()

    # Clear existing content (naive approach: just clear element tree or start fresh)
    # A cleaner way is to just create a new document instance if we want to replace everything
    # doc = Document()  <-- REMOVED: This was causing the overwrite. Now we maintain the existing doc structure.

    # Style Configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    print("Writing content...")
    
    code_block = False
    
    for line in lines:
        stripped = line.strip()
        
        # Handle Headers
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        
        # Handle Code Blocks
        elif stripped.startswith('```'):
            code_block = not code_block
            continue
            
        # Handle Lists
        elif stripped.startswith('* ') or stripped.startswith('- '):
            p = doc.add_paragraph(stripped[2:], style='List Bullet')
        
        # Handle Tables (Basic Text dump for now as complex table parsing is hard without a library)
        elif stripped.startswith('|'):
            p = doc.add_paragraph(stripped)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            font = p.add_run().font
            font.name = 'Courier New'
            font.size = Pt(9)

        # Standard Text
        else:
            if stripped:
                if code_block:
                    p = doc.add_paragraph(stripped)
                    p.style = 'No Spacing'
                    run = p.runs[0]
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 50, 0)
                else:
                    doc.add_paragraph(stripped)

    print(f"Saving to: {DOCX_TARGET}")
    doc.save(DOCX_TARGET)
    print("Optimization Complete.")

if __name__ == "__main__":
    update_docx()
